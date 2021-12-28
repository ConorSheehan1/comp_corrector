"""
These tests are considered integration tests because they have side effects.
They write to the file system (in tests/fixtures/example).
They rely on gcc being available.
"""

# Standard Library
# TODO: stub file system?
import os
import platform
import re
import shutil

# TODO: rename to test_ui and setup mocking with tkinter
# for now, run all the steps that ui.py would run in main, but without errors and warnings for the user
import unittest

# Third party
from docx import Document

from src.file_management.compile import compile_c
from src.file_management.feedback import create_feedback_file, get_missing_names

# functions under test
from src.file_management.zip_archives import setup_safe_mode, unzip, unzip_outer


class Base:
    # needed as class variables to run teardown after all tests, instead of each test.
    cwd = os.path.join("tests", "fixtures")
    safe_cwd = os.path.join(cwd, "example")

    def __init__(self, methodName="runTest"):
        unittest.TestCase.__init__(self, methodName)
        self.safe_cwd = self.__class__.safe_cwd
        self.zip_path = os.path.join(self.cwd, "example.zip")
        self.example_feedback_file = os.path.join(self.cwd, "example_feedback.docx")
        self.safe_zip_path = os.path.join(self.safe_cwd, "example.zip")
        self.example_student_dir = os.path.join(self.safe_cwd, "final_fake_student_2012347")
        self.example_student_code = os.path.join(
            self.example_student_dir, "assignment1", "assignment1.c"
        )
        self.feedback_file = os.path.join(self.safe_cwd, "feedback.docx")
        self.all_names = ["fake_student", "other_fake_student", "final_fake_student"]
        self.missing_names = ["missing_student"]

    @classmethod
    def tearDownClass(cls):
        # delete non-empty example dir
        shutil.rmtree(cls.safe_cwd, ignore_errors=True)

    # https://stackoverflow.com/a/18627017/6305204
    # nosetests runs tests in order by name
    # not ideal, but need to enforce order until testing tkinter directly, or using different UI framework.
    def test_01_safe_mode(self):
        """
        setup_safe_mode should create a new dir and move the main .zip archive into it before extracting
        """
        # example dir should not exist yet
        assert not os.path.exists(self.safe_cwd)
        assert not os.path.exists(self.safe_zip_path)
        setup_safe_mode(self.cwd, self.zip_path)
        assert os.path.exists(self.safe_cwd)
        assert os.path.exists(self.zip_path)


class TestMainSafeModeSharedState(Base, unittest.TestCase):
    def test_02_unzip(self):
        """
        unzip_outer should unzip the main .zip archive
        """
        assert not os.path.exists(self.example_student_dir)
        # when names is [""], everything is extracted
        unzip_outer(self.safe_zip_path, [""])
        assert os.path.exists(self.example_student_dir)

    def test_03_unzip_inner(self):
        """
        unzip should unzip each students .zip archives nested in the main zip archive
        """
        assert not os.path.exists(self.example_student_code)
        unzip(self.safe_cwd)
        assert os.path.exists(self.example_student_code)

    def test_04_missing_names(self):
        """
        get_missing_names should return names of students who didn't submit anything
        """
        missing_names = get_missing_names(self.safe_cwd, self.all_names + self.missing_names)
        self.assertListEqual(self.missing_names, missing_names)

    def test_05_compile(self):
        """
        compile_c should compile assignment1.c to assignment1 executable.
        it should also return the error found compiling other_fake_student_2012346/assignment1.c
        """
        # compiled_file = self.example_student_code.replace(".c", "")
        compiled_file = re.sub(r"\.c$", "", self.example_student_code)

        # gcc output always add .exe on windows
        if platform.system() == "Windows":
            compiled_file = f"{compiled_file}.exe"

        assert not os.path.exists(compiled_file)
        # hide output for tests
        errors = compile_c(self.safe_cwd, capture_output=True)
        assert errors == 1
        assert os.path.exists(compiled_file)

    def test_06_feedback_exists(self):
        """
        create_feedback_file should create feedback.docx
        """
        assert not os.path.exists(self.feedback_file)
        create_feedback_file(self.safe_cwd, self.all_names, self.missing_names)
        assert os.path.exists(self.feedback_file)

    def test_07_feedback_content(self):
        expected_table = Document(self.example_feedback_file).tables[0]
        actual_table = Document(self.feedback_file).tables[0]

        for expected_row, actual_row in zip(expected_table.rows, actual_table.rows):
            for expected_cell, actual_cell in zip(expected_row.cells, actual_row.cells):
                assert expected_cell.text == actual_cell.text


class TestMainSafeModeSpecificStudents(Base, unittest.TestCase):
    def test_02_unzip(self):
        """
        when students are passed, unzip_outer should only extract files containing that students name
        """
        assert not os.path.exists(self.example_student_dir)
        unzip_outer(self.safe_zip_path, ["fake_student", "other_fake_student"])

        # final_fake_student is the only one not extracted, all others are extracted
        assert not os.path.exists(self.example_student_dir)
        for dirname in ["fake_student_2012345", "other_fake_student_2012346"]:
            student_dir = os.path.join(self.safe_cwd, dirname)
            assert os.path.exists(student_dir)


if __name__ == "__main__":
    unittest.main()
